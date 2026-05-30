import os
from typing import Optional


class ExportService:
    """报告导出服务：Markdown → PDF / DOCX"""

    @classmethod
    def _is_table_separator(cls, line: str) -> bool:
        stripped = line.replace(" ", "").replace("|", "").replace("-", "").replace(":", "")
        return stripped == "" and "---" in line.replace(" ", "")

    @classmethod
    def _parse_table_row(cls, line: str) -> list[str]:
        cells = line.strip().strip("|").split("|")
        return [c.strip() for c in cells]

    @classmethod
    def _group_lines(cls, lines: list[str]) -> list:
        blocks = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith("|") and " | " in line and i + 1 < len(lines) and cls._is_table_separator(lines[i + 1]):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                blocks.append(("table", table_lines))
            else:
                blocks.append(("line", line))
                i += 1
        return blocks

    @classmethod
    def _get_cjk_font(cls) -> tuple[Optional[str], Optional[str]]:
        """查找系统中可用的 CJK 字体，返回 (字体路径, 字体名)"""
        font_candidates = [
            # Windows 常见中文字体
            (os.path.join(os.environ.get("WINDIR", "C:/Windows"), "Fonts", "msyh.ttc"), "MicrosoftYaHei"),
            (os.path.join(os.environ.get("WINDIR", "C:/Windows"), "Fonts", "simsun.ttc"), "SimSun"),
            (os.path.join(os.environ.get("WINDIR", "C:/Windows"), "Fonts", "simhei.ttf"), "SimHei"),
            (os.path.join(os.environ.get("WINDIR", "C:/Windows"), "Fonts", "msyh.ttf"), "MicrosoftYaHei"),
            # Linux/macOS 常见中文字体
            ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "WQY"),
            ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
        ]
        for path, name in font_candidates:
            if os.path.isfile(path):
                return path, name
        return None, None

    @classmethod
    def export_pdf(cls, markdown_content: str, output_path: str) -> Optional[str]:
        """将 Markdown 转为 PDF（使用 fpdf2，纯 Python 实现）"""
        try:
            from fpdf import FPDF

            font_path, font_name = cls._get_cjk_font()
            if font_path:
                use_cjk = True
            else:
                use_cjk = False
                font_name = "Helvetica"

            class PDF(FPDF):
                def header(self):
                    if use_cjk:
                        self.set_font(font_name, "", 10)
                    else:
                        self.set_font("Helvetica", "", 10)
                    self.set_text_color(128, 128, 128)
                    self.cell(0, 10, "AI Generated Report", align="R", new_x="LMARGIN", new_y="NEXT")
                    self.ln(2)

                def footer(self):
                    self.set_y(-15)
                    if use_cjk:
                        self.set_font(font_name, "", 9)
                    else:
                        self.set_font("Helvetica", "", 9)
                    self.set_text_color(128, 128, 128)
                    self.cell(0, 10, f"Page {self.page_no()}", align="C")

            pdf = PDF()
            if use_cjk:
                pdf.add_font(font_name, "", font_path, uni=True)
                pdf.add_font(font_name, "B", font_path, uni=True)
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            lines = markdown_content.split("\n")
            blocks = cls._group_lines(lines)

            for block_type, block_data in blocks:
                if block_type == "table":
                    rows_data = []
                    for tl in block_data:
                        if cls._is_table_separator(tl):
                            continue
                        rows_data.append(cls._parse_table_row(tl))
                    if not rows_data:
                        continue
                    num_cols = max(len(r) for r in rows_data)
                    col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / max(num_cols, 1)
                    if use_cjk:
                        pdf.set_font(font_name, "", 9)
                    else:
                        pdf.set_font("Helvetica", "", 9)
                    for ri, row_cells in enumerate(rows_data):
                        max_h = 6
                        x_start = pdf.get_x()
                        y_start = pdf.get_y()
                        for ci in range(num_cols):
                            cell_text = row_cells[ci] if ci < len(row_cells) else ""
                            if ri == 0:
                                if use_cjk:
                                    pdf.set_font(font_name, "B", 9)
                                else:
                                    pdf.set_font("Helvetica", "B", 9)
                            else:
                                if use_cjk:
                                    pdf.set_font(font_name, "", 9)
                                else:
                                    pdf.set_font("Helvetica", "", 9)
                            pdf.set_xy(x_start + ci * col_w, y_start)
                            pdf.multi_cell(col_w, 5, cell_text, border=1)
                            max_h = max(max_h, pdf.get_y() - y_start)
                        pdf.set_y(y_start + max_h)
                    pdf.ln(3)
                    continue
                line = block_data
                if not line:
                    pdf.ln(4)
                    continue
                try:
                    if line.startswith("# "):
                        if use_cjk:
                            pdf.set_font(font_name, "B", 16)
                        else:
                            pdf.set_font("Helvetica", "B", 16)
                        pdf.set_text_color(0, 102, 204)
                        pdf.cell(0, 10, line[2:], new_x="LMARGIN", new_y="NEXT")
                        pdf.ln(2)
                    elif line.startswith("## "):
                        if use_cjk:
                            pdf.set_font(font_name, "B", 14)
                        else:
                            pdf.set_font("Helvetica", "B", 14)
                        pdf.set_text_color(51, 51, 51)
                        pdf.cell(0, 8, line[3:], new_x="LMARGIN", new_y="NEXT")
                        pdf.ln(1)
                    elif line.startswith("### "):
                        if use_cjk:
                            pdf.set_font(font_name, "B", 12)
                        else:
                            pdf.set_font("Helvetica", "B", 12)
                        pdf.set_text_color(85, 85, 85)
                        pdf.cell(0, 7, line[4:], new_x="LMARGIN", new_y="NEXT")
                    elif line.startswith("- "):
                        pdf.set_text_color(0, 0, 0)
                        bullet_text = chr(149) + " " + line[2:]
                        cls._write_inline_pdf(pdf, bullet_text, font_name if use_cjk else "Helvetica", 11, 6)
                    else:
                        pdf.set_text_color(0, 0, 0)
                        cls._write_inline_pdf(pdf, line, font_name if use_cjk else "Helvetica", 11, 6)
                except Exception as line_err:
                    # Skip lines that can't be rendered
                    import logging
                    logging.getLogger(__name__).debug(f"PDF line render skipped: {line_err}, line: {line[:50]}")
                    pdf.set_x(pdf.l_margin)
                    continue

            pdf.output(output_path)
            return output_path
        except Exception as e:
            import logging
            logging.getLogger(__name__).warning(f"PDF export failed: {e}")
            return None

    @classmethod
    def _parse_inline(cls, text: str):
        """解析内联 Markdown，返回 [(style, text), ...]"""
        parts = []
        i = 0
        while i < len(text):
            if text.startswith('**', i):
                end = text.find('**', i + 2)
                if end == -1:
                    i += 2
                    continue
                if end > i + 2:
                    parts.append(('bold', text[i + 2:end]))
                i = end + 2
            elif text[i] == '*':
                end = text.find('*', i + 1)
                if end == -1:
                    i += 1
                    continue
                if end > i + 1:
                    parts.append(('italic', text[i + 1:end]))
                i = end + 1
            else:
                next_bold = text.find('**', i)
                next_italic = text.find('*', i)
                candidates = [pos for pos in (next_bold, next_italic) if pos != -1]
                next_marker = min(candidates) if candidates else len(text)
                parts.append(('normal', text[i:next_marker]))
                i = next_marker
        return parts

    @classmethod
    def _write_inline_pdf(cls, pdf, text: str, font_name: str, size: int, line_height: float):
        """用 fpdf2 write() 渲染带行内 Markdown 格式的文本"""
        parts = cls._parse_inline(text)
        for style, part_text in parts:
            if style == 'bold':
                pdf.set_font(font_name, "B", size)
            elif style == 'italic':
                pdf.set_font(font_name, "I", size)
            else:
                pdf.set_font(font_name, "", size)
            # 超出页面右边界时换行
            available = pdf.w - pdf.r_margin - pdf.get_x()
            width = pdf.get_string_width(part_text)
            if width > available and pdf.get_x() > pdf.l_margin:
                pdf.ln(line_height)
            pdf.write(line_height, part_text)
        pdf.ln(line_height)

    @classmethod
    def _set_docx_style_font(cls, style, font_name: str, font_size=None):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        style.font.name = font_name
        if font_size:
            style.font.size = font_size
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)

    @classmethod
    def export_docx(cls, markdown_content: str, output_path: str) -> Optional[str]:
        """将 Markdown 转为 DOCX"""
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()

            _, cjk_font = cls._get_cjk_font()
            font_name = cjk_font or 'Calibri'
            size_map = {
                'Normal': Pt(11),
                'Heading 1': Pt(18),
                'Heading 2': Pt(14),
                'Heading 3': Pt(12),
                'List Bullet': Pt(11),
                'List Number': Pt(11),
            }
            for name, size in size_map.items():
                try:
                    cls._set_docx_style_font(doc.styles[name], font_name, size)
                except KeyError:
                    pass

            def _add_paragraph(text, style_name='Normal'):
                p = doc.add_paragraph(style=style_name)
                parts = cls._parse_inline(text)
                for style, part_text in parts:
                    run = p.add_run(part_text)
                    run.font.name = font_name
                    run.font.size = size_map.get(style_name, Pt(11))
                    if style == 'bold':
                        run.bold = True
                    elif style == 'italic':
                        run.italic = True
                    from docx.oxml.ns import qn
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                return p

            lines = markdown_content.split("\n")
            blocks = cls._group_lines(lines)

            for block_type, block_data in blocks:
                if block_type == "table":
                    rows_data = []
                    for tl in block_data:
                        if cls._is_table_separator(tl):
                            continue
                        rows_data.append(cls._parse_table_row(tl))
                    if not rows_data:
                        continue
                    num_cols = max(len(r) for r in rows_data)
                    table = doc.add_table(rows=len(rows_data), cols=num_cols)
                    table.style = 'Table Grid'
                    for ri, row_cells in enumerate(rows_data):
                        for ci in range(num_cols):
                            cell = table.cell(ri, ci)
                            cell_text = row_cells[ci] if ci < len(row_cells) else ""
                            cell.text = ""
                            p = cell.paragraphs[0]
                            parts = cls._parse_inline(cell_text)
                            for style, part_text in parts:
                                run = p.add_run(part_text)
                                run.font.name = font_name
                                run.font.size = Pt(10)
                                if style == 'bold' or ri == 0:
                                    run.bold = True
                                elif style == 'italic':
                                    run.italic = True
                                from docx.oxml.ns import qn
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    doc.add_paragraph("")
                else:
                    line = block_data
                    if not line:
                        continue
                    if line.startswith("# "):
                        _add_paragraph(line[2:], 'Heading 1')
                    elif line.startswith("## "):
                        _add_paragraph(line[3:], 'Heading 2')
                    elif line.startswith("### "):
                        _add_paragraph(line[4:], 'Heading 3')
                    elif line.startswith("- "):
                        _add_paragraph(line[2:], 'List Bullet')
                    elif line[:3].replace(".", "").isdigit() and ". " in line[:5]:
                        _add_paragraph(line.split(". ", 1)[1], 'List Number')
                    else:
                        _add_paragraph(line)

            doc.save(output_path)
            return output_path
        except ImportError:
            return None
        except Exception:
            return None
