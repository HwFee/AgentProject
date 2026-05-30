from pathlib import Path
from zipfile import ZipFile

from docx import Document

from services.export import ExportService


def test_docx_export_removes_markdown_bold_markers_and_keeps_size(tmp_path: Path):
    output_path = tmp_path / "report.docx"
    markdown = (
        '→ **"预防式降息"场景的脆弱性**：如果美国经济实现"软着陆"，'
        "美联储将缓慢、小幅降息。 **核心风险点在于**：市场低估了就业市场"
        "突然崩溃的可能性。 **若紧缩政策**的滞后效应在 Q4 集中显现。"
    )

    assert ExportService.export_docx(markdown, str(output_path)) == str(output_path)

    with ZipFile(output_path) as docx:
        document_xml = docx.read("word/document.xml").decode("utf-8")
    assert "**" not in document_xml

    paragraph = Document(output_path).paragraphs[0]
    assert paragraph.text.count("*") == 0
    assert {run.font.size for run in paragraph.runs} == {paragraph.runs[0].font.size}
    assert any(run.bold and "预防式降息" in run.text for run in paragraph.runs)
    assert any(run.bold and "核心风险点在于" in run.text for run in paragraph.runs)


def test_docx_export_drops_unmatched_markdown_markers(tmp_path: Path):
    output_path = tmp_path / "report.docx"
    markdown = "正文 **残缺粗体标记 不应把星号导出到 Word，后续文字保持可读"

    assert ExportService.export_docx(markdown, str(output_path)) == str(output_path)

    paragraph = Document(output_path).paragraphs[0]
    assert paragraph.text == "正文 残缺粗体标记 不应把星号导出到 Word，后续文字保持可读"
